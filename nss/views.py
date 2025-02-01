from django.shortcuts import render, get_object_or_404, redirect,reverse
from django.http import HttpResponse,HttpResponseBadRequest
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from .models import *
from django.contrib.auth.models import Group,User
from django.db import connection
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from .decorators import group_required 
from datetime import datetime
from .filters import *
from .forms import *
import os
from django.template.loader import get_template
from django.utils.timezone import now
from django.conf import settings
from docx import Document
from docx.shared import Inches, Cm

from django.contrib.auth.forms import PasswordChangeForm
from django.contrib.auth import update_session_auth_hash



@login_required()
def change_password(request):
    if request.method == 'POST':
        form = PasswordChangeForm(request.user, request.POST)
        if form.is_valid():
            user = form.save()
            update_session_auth_hash(request, user)  # Important! Keeps the user logged in after password change
            messages.success(request, 'Your password was successfully updated!')
            return redirect('change_password')
        else:
            messages.error(request, 'Please correct the error below.')
    else:
        form = PasswordChangeForm(request.user)
    return render(request, 'registration/change_password.html', {'form': form})



@login_required()
@group_required('po','vs','admin')
def ns(request):
    try:
        events = Event.objects.filter(date__lte=now()).order_by('-date')[:5]
        volunteers = volunteer.objects.all()
        pending_approvals = Attendance_status.objects.filter(status="pending for approval")
        
        context = {
            'events': events,
            'volunteers': volunteers,
            'pending_approvals': pending_approvals
        }
        return render(request, 'nss/home.html', context)
    except Exception:
        return render(request,'nss/error.html')
    


@login_required()
@group_required('po','vs')
def add_volunteer(request):
    try:
        if request.method == "POST":
            form = addVolunteerForm(request.POST, request.FILES)
            if form.is_valid():
                form.save()
                return redirect('view_volunteer')
        else:
            form = addVolunteerForm()

        return render(request, 'nss/add_volunteer.html', {'form': form})
    except Exception as e:
        print(str(e)) 
        return render(request, 'nss/error.html')

@login_required()
@group_required('po','vs')
def edit_volunteer(request, pk):
    try:
        volunteer_instance = get_object_or_404(volunteer, pk=pk)
        
        if request.method == "POST":
            form = addVolunteerForm(request.POST, request.FILES, instance=volunteer_instance)
            if form.is_valid():
                form.save()  
                return redirect('view_volunteer')  
        else:
            form = addVolunteerForm(instance=volunteer_instance)

        context = {
            'form': form,
            'volunteer': volunteer_instance,
        }
        
        return render(request, 'nss/edit_volunteer.html', context)
    
    except Exception as e:
        print(str(e))
        return render(request, 'nss/error.html')

    
@login_required()
@group_required('po','vs')
def view_volunteer(request):
    try:
        volunteer_list = volunteer.objects.filter(status='active')
        volunteer_filter = VolunteerFilter(request.GET, queryset=volunteer_list)
        
        items_per_page = request.GET.get('items_per_page', 10)
        
        try:
            items_per_page = int(items_per_page)
        except ValueError:
            items_per_page = 10
        
        paginator = Paginator(volunteer_filter.qs, items_per_page)
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)
        
        return render(request, 'nss/view_volunteer.html', {'filter': volunteer_filter, 'page_obj': page_obj, 'items_per_page': items_per_page})
    except Exception:
        return render(request,'nss/error.html')

@login_required()
@group_required('po','vs')
def view_attendance2(request):
    try:
        volunteer_list = volunteer.objects.all()
        volunteer_filter = VolunteerFilter2(request.GET, queryset=volunteer_list)
        
        items_per_page = request.GET.get('items_per_page', 10)
        
        try:
            items_per_page = int(items_per_page)
        except ValueError:
            items_per_page = 10
        
        paginator = Paginator(volunteer_filter.qs, items_per_page)
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)
        
        return render(request, 'nss/full_attendance.html', {'filter': volunteer_filter, 'page_obj': page_obj, 'items_per_page': items_per_page})
    except Exception:
        return render(request,'nss/error.html')
@login_required()
@group_required('po','vs')
def view_attendance(request):
    try:
        eve = {
            'even': Attendance_status.objects.all().order_by('date').select_related('event')
        }
        
        if request.method == "POST":
            at_status_id = request.POST.get('event')
            return redirect('view_attendance3', status_id=at_status_id)
            
        return render(request, 'nss/view_attendance.html', eve)
    except Exception:
        return render(request,'nss/error.html')
@login_required()
@group_required('po','vs')
def volunteer_details(request, volunteer_name):
    try:
        # Assuming you have a Volunteer model with a 'name' field
        vol={
            'voluntee':volunteer.objects.filter(volunteer_id=volunteer_name)
        }
        
        ev={
            'even':Attendance.objects.filter(volunteer=volunteer_name)
        }
        # Pass the volunteer details to the template
        return render(request, 'nss/volunteer_details.html', {**vol,**ev})
    except Exception:
        return render(request,'nss/error.html')
@login_required()
@group_required('po','vs')
def event(request):
    try:

        return render(request,'nss/event.html')

    except Exception:
        return render(request,'nss/error.html')


@login_required()
@group_required('po','vs')
def add_event(request):
    try:
        if request.method=="POST":
            event_name=request.POST.get('event_name')
            date=request.POST.get('date')
            ev=Event(event_name=event_name,date=date)
            ev.save()
            message="Submitted Successfully"
            return redirect(reverse('add_event') + '?message=' + message)
        return render(request,'nss/add_event.html')
    except Exception:
        return render(request,'nss/error.html')


@login_required()
@group_required('po','vs')
def event_details(request):

        eve = {
            'even': Event.objects.all().order_by('date').values()
        }
        if request.method=="POST":
            event_id=request.POST.get('event_name')
            des=request.POST.get('des')
            expense=request.POST.get('expense')
            event_id=Event.objects.get(event_id=event_id)
            event_exists = Event_details.objects.filter(event=event_id).exists()
            if not event_exists:
                ev=Event_details(event=event_id,des=des,expense=expense)
                ev.save()
                message="Submitted Successfully"
                return redirect(reverse('event_details') + '?message1=' + message)
            else:
                ev=Event_details.objects.get(event=event_id)
                ev.des=des
                ev.expense=expense
                ev.save()
                message="Submitted Successfully"
                return redirect(reverse('event_details') + '?message1=' + message)
        return render(request,'nss/event_details.html',eve)

@login_required()
@group_required('po','vs')
def report(request):
    try:
        report={
            'event':Event.objects.all()
        }
        pics={
            'pics':Event_Photos.objects.all()
        }
        details={
            'details':Event_details.objects.all()
        }
        return render(request,'nss/report.html',{**report,**pics,**details})
    except Exception:
        return render(request,'nss/error.html')
@login_required()
@group_required('po','vs')
def event_photos(request):
    try:
        eve = {
            'even': Event.objects.all().order_by('date').values()
        }
        if request.method == "POST" and request.FILES:
            photo=request.FILES.get('photo')
            event_id=request.POST.get('event_name')
            event_id=Event.objects.get(event_id=event_id)
            if Event_Photos.objects.filter(event=event_id).count() >= 3:
                message = "Three images are already uploaded for this event."
                return redirect(reverse('event_photos') + '?message1=' + message)
            else:
                ev=Event_Photos(photo=photo,event=event_id)
                ev.save()
                message = "Uploaded Sucessfully"
                return redirect(reverse('event_photos') + '?message2=' + message)
        return render(request,'nss/add_photos.html',eve)
    except Exception:
        return render(request,'nss/error.html')
@login_required()
@group_required('po','vs')
def event2(request):
    try:
        return render(request,'nss/event2.html')
    except Exception:
        return render(request,'nss/error.html')


@login_required()
@group_required('po','vs')
def delete2(request, pk):
    try:
        vol={
            'voluntee':volunteer.objects.filter(volunteer_id=pk)
        }
        
        ev={
            'even':Attendance.objects.filter(volunteer=pk)
        }
        # Pass the volunteer details to the template
        return render(request, 'nss/delete_volunteer.html', {**vol,**ev})
    except Exception:
        return render(request,'nss/error.html')
@login_required()
@group_required('po','vs')
def delete_volunteer(request, pk):
    try:
        Volunteer = get_object_or_404(volunteer, pk=pk)
        Volunteer.delete()
        return redirect('view_volunteer') 
    except Exception:
        return render(request,'nss/error.html')
@login_required()
@group_required('po','vs')
def report_list(request):
    try:
        report={
            'event':Event.objects.all().order_by('date')
        }
        pics={
            'pics':Event_Photos.objects.all()
        }
        details={
            'details':Event_details.objects.all()
        }
        return render(request,'nss/report_list.html',{**report,**pics,**details})
    except Exception:
        return render(request,'nss/error.html')
@login_required()
@group_required('po','vs')
def report_list_more(request,pk):

    event={
        'eve':Event.objects.filter(event_id=pk)
    }
    pics={
        'pics':Event_Photos.objects.filter(event_id=pk)
    }
    desc={
        'desc':Event_details.objects.filter(event_id=pk)
    }
    return render(request,'nss/report_list_more.html',{**event,**pics,**desc})

@login_required()
@group_required('po','vs')
def view_event(request):
    try:
        event={
            'eve':Event.objects.all()
        }
        return render(request,'nss/view_event.html',event)
    except Exception:
        return render(request,'nss/error.html')

@login_required()
@group_required('po','vs')
def delete_event(request,pk):
    try:
        event = get_object_or_404(Event, pk=pk)
        print(event)
        event.delete()
        return redirect('view_event')
    except Exception:
        return render(request,'nss/error.html')
@login_required()
@group_required('po')
def approve_attendance(request,pk):
    try:
        att=Attendance_status.objects.get(status_id=pk)
        att.status="approved"
        att.save()
        return redirect('view_attendance3', status_id=pk) 
    except Exception:
        return render(request,'nss/error.html')
@login_required()
@group_required('po')
def reject_attendance(request,pk):
    try:
        att=Attendance_status.objects.get(status_id=pk)
        att.status="rejected"
        att.save()
        return redirect('view_attendance3', status_id=pk) 
    except Exception:
        return render(request,'nss/error.html')


@login_required()
@group_required('po','vs')
def add_attendance(request):
    try:
        eve = {
            'even': Event.objects.all().order_by('date').values()
        }
        return render(request, 'nss/add_attendance.html',eve)
    except Exception:
        return render(request,'nss/error.html')

@login_required()
@group_required('po','vs')
def dep_wise(request):
    try:
        if request.method=="POST":
            unit=request.POST.get('unit')
            event=request.POST.get('event')
            event=Event.objects.get(event_id=event)
            dept={
                'pog':Programme.objects.all(),
                'unit':unit,
                'event':event,
                'vol':volunteer.objects.filter(unit=unit,status='active'),
                'list':[1,2,3]
                }
            return render(request,'nss/unit_wise.html',dept)
    except Exception:
        return render(request,'nss/error.html')

@login_required()
@group_required('po','vs')
def attendance(request,pk,unit):
    try:
        if request.method == "POST":
            date = request.POST.get('date')
            event=Event.objects.get(event_id=pk)

            event = get_object_or_404(Event, event_id=pk)
            
            # Check if attendance for this event on the given date already exists
            existing_attendance = Attendance_status.objects.filter(event=event, date=date,unit=unit).exists()
            
            if existing_attendance:
                status_id = Attendance_status.objects.get(event=event, date=date, unit=unit).status_id
                messages.error(request,f"Attendance for the event {event} on the date {date} already exists.")
            else:
                id_list = request.POST.getlist('volunteers')
                time=request.POST.get('time')
                at_status=Attendance_status(event=event,date=date,unit=unit)
                at_status.save()
                for volunteer_id in id_list:
                    volunteer_instance= volunteer.objects.get(volunteer_id=volunteer_id)
                    
                    # Save the attendance record
                    event=Event.objects.get(event_id=pk)
                    att = Attendance(Attendance_status=at_status,volunteer=volunteer_instance,date=date,event=event,no_of_hours=time)
                    print(att)           
                    att.save()

        return redirect('view_attendance')
        
    except Exception:
        return render(request,'nss/error.html')

@group_required('po')
def promote(request):
    try:
        objec=volunteer.objects.all()
        for i in objec:
            if i.year != 3:
                i.year+=1
            else:
                i.status='inactive'
            i.save()
        message="Promoted"
        return redirect(reverse('promote_check') +'?message=' + message)
    except Exception:
        return render(request,'nss/error.html')

@login_required
@group_required('po','vs')
def more_attendance(request,pk):
    try:
        vol={
            'vol':volunteer.objects.filter(volunteer_id=pk)
        }
        return render(request,'nss/more_attendance.html',vol)
    except Exception:
        return render(request,'nss/error.html')

@login_required()
@group_required('po','vs')
def delete_attendance(request,pk):
    try:
        att=get_object_or_404(Attendance_status,status_id=pk)
        att.delete()
        return redirect('view_attendance')
    except Exception:
        return render(request,'nss/error.html')
@login_required()
@group_required('po','vs')
def delete_images(request,pk,ev):
    try:
        pic=get_object_or_404(Event_Photos,id=pk)
        pic.delete()
        photo_path = os.path.join(settings.MEDIA_ROOT, str(pic.photo))
        if os.path.exists(photo_path):
            os.remove(photo_path) 
        message = "Deleted Successfully."
        return redirect(reverse('edit_event', args=[ev])+ '?message=' + message) 
    except Exception:
        return render(request,'nss/error.html')
@login_required()
@group_required('po','vs')
def edit_event(request,pk):

    eve={
        'eve':Event.objects.filter(event_id=pk)
    }
    ev=Event.objects.get(event_id=pk)
    event_Photos = Event_Photos.objects.filter(event=ev)
    event_details, created = Event_details.objects.get_or_create(event=ev)
    event1 = get_object_or_404(Event, pk=pk)
    if request.method=='POST':
        event_name=request.POST.get('event_name')
        date=request.POST.get('date')
        des=request.POST.get('des')
        event_details.des=des
        event1.event_name=event_name
        event1.date=date
        if event_details.des:
            event_details.save()
        event1.save()
        return redirect('view_event')
    return render(request,'nss/edit_event.html',eve)

@login_required()
@group_required('po','vs')
def promote_check(request):
    try:
        return render(request,'nss/promote.html')
    except Exception:
        return render(request,'nss/error.html')

@login_required()
@group_required('po','vs')
def monthly_report(request):
    try:
        if request.method=="POST":
            year=request.POST.get('year')
            print(year)
            month=request.POST.get('month')
            print(month)
            events = Event.objects.filter(date__year=year, date__month=month)
            details = Event_details.objects.filter(event__in=events)
            pics = Event_Photos.objects.filter(event__in=events)
            return render(request, 'nss/report.html', {'event': events, 'details': details, 'pics': pics})
    except Exception:
        return render(request,'nss/error.html')
@login_required()
@group_required('po','vs')
def yearly_report(request):
    try:
        if request.method=='POST':
            fromyear=request.POST.get('fromyear')
            toyear=request.POST.get('toyear')
            events = Event.objects.filter(date__gte=fromyear, date__lte=toyear)
            print(events)  # Add a debug print statement
            details = Event_details.objects.filter(event__in=events)
            pics = Event_Photos.objects.filter(event__in=events)
            return render(request, 'nss/report.html', {'event': events, 'details': details, 'pics': pics})
    except Exception:
        return render(request,'nss/error.html')
@login_required()
@group_required('po','vs')
def select_month(request):
    try:
        return render(request,'nss/select_month.html')
    except Exception:
        return render(request,'nss/error.html')
@login_required()
@group_required('po','vs')
def select_year(request):
    try:
        return render(request,'nss/select_year.html')
    except Exception:
        return render(request,'nss/error.html')


@login_required()
@group_required('po','vs')
def view_attendance3(request,status_id):
    try:
        at_status=Attendance_status.objects.get(status_id=status_id)
        res = {
            'resul': Attendance.objects.filter(Attendance_status=at_status).order_by('date').select_related('volunteer')
        }
        status={
            'status':Attendance_status.objects.get(status_id=status_id)
        }
        return render(request, 'nss/view_attendance3.html',{**status,**res,'status_id': status_id})
    except Exception:
        return render(request,'nss/error.html')

# ADMIN VIEWS

@login_required
@group_required('admin','po')
def manage_users(request):
    try:
        groups = Group.objects.all()
        
        if request.method == 'POST':
            if request.user.groups.filter(name='admin').exists():
                form = UserForm(request.POST)
            else:
                form=UserForm3(request.POST)
            if form.is_valid():
                form.save()
                return redirect('manage_users')
        else:
            if request.user.groups.filter(name='admin').exists():
                form = UserForm()
                users = User.objects.all()
            else:
                group_vs = Group.objects.get(name='vs')
                users = User.objects.filter(groups=group_vs)
                form=UserForm3
        return render(request, 'admin/manage_users.html', {'groups': groups,'users': users, 'form': form})
    except Exception:
        return render(request,'nss/error.html')
@login_required()
@group_required('admin')
def manage_groups(request):
    try:
        groups = Group.objects.all()
        if request.method == 'POST':
            form = GroupForm(request.POST)
            if form.is_valid():
                form.save()
                return redirect('manage_groups')
        else:
            form = GroupForm()
        return render(request, 'admin/manage_groups.html', {'groups': groups, 'form': form})

    except Exception:
        return render(request,'nss/error.html')
@login_required()
@group_required('admin','po')
def edit_user(request, pk):
    try:
        user = get_object_or_404(User, pk=pk)
        if request.method == 'POST':
            if request.user.groups.filter(name='admin').exists():
                form = UserForm2(request.POST, instance=user)
            else:
                form = UserForm4(request.POST, instance=user)
            if form.is_valid():
                form.save()
                return redirect('manage_users')
        else:
            if request.user.groups.filter(name='admin').exists():
                form = UserForm2(instance=user)
            else:
                form=UserForm4(instance=user)
        return render(request, 'admin/edit_user.html', {'form': form})
    except Exception:
        return render(request,'nss/error.html')
@login_required()
@group_required('admin')
def edit_group(request, pk):
    try:
        group = get_object_or_404(Group, pk=pk)
        if request.method == 'POST':
            form = GroupForm(request.POST, instance=group)
            if form.is_valid():
                form.save()
                return redirect('manage_groups')
        else:
            form = GroupForm(instance=group)
        return render(request, 'admin/edit_group.html', {'form': form})
    except Exception:
        return render(request,'nss/error.html')

@login_required()
@group_required('admin','po')
def delete_user(request, pk):
    try:
        user = get_object_or_404(User, pk=pk)
        if request.method == 'POST':
            user.delete()
            return redirect('manage_users')
        return render(request, 'admin/confirm_delete.html', {'object': user})
    except Exception:
        return render(request,'nss/error.html')

@login_required()
@group_required('admin')
def delete_group(request, pk):
    try:
        group = get_object_or_404(Group, pk=pk)
        if request.method == 'POST':
            group.delete()
            return redirect('manage_groups')
        return render(request, 'admin/confirm_delete.html', {'object': group})
    except Exception:
        return render(request,'nss/error.html')

@login_required()
@group_required('po','vs')
def camp(request):
    
    return render(request,'camp/camp.html')

@login_required()
@group_required('po','vs')
def addcamp(request):
    if request.method=="POST":
        camp_name=request.POST.get('camp_name')
        fromdate=request.POST.get('fromdate')
        todate=request.POST.get('todate')
        ca=Camp(camp_name=camp_name,fromdate=fromdate,todate=todate)
        ca.save()
        message="Submitted Successfully"
        return redirect(reverse('add_camp') + '?message=' + message)
    return render(request,'camp/add_camp.html')

@login_required()
@group_required('po','vs')
def campattendance(request):
    dept={
        'pog':Programme.objects.all(),
        'vol':volunteer.objects.filter(status='active'),
        'list':[2],
        'camps':Camp.objects.all()
       }
    return render(request,'camp/campattendance.html',dept)


@login_required()
@group_required('po','vs')
def add_camp_attendance(request):
    if request.method == "POST":
        camp = request.POST.get('camp')
        camp=Camp.objects.get(camp_id=camp)
        # Check if attendance for this event on the given date already exists
        existing_attendance = Camp_Attendance.objects.filter(camp=camp).exists()
        
        if existing_attendance:
            message = "Attendance Exists"
            return redirect(reverse('campattendance') + '?message1=' + message)
        else:
            id_list = request.POST.getlist('volunteers')
            for volunteer_id in id_list:
                volunteer_instance= volunteer.objects.get(volunteer_id=volunteer_id)
                # Save the attendance record
                att = Camp_Attendance(camp=camp,volunteer=volunteer_instance)
                print(att)           
                att.save()
            message = "Attendance Added"
            return redirect(reverse('campattendance') + '?message2=' + message)

@login_required()
@group_required('po','vs')
def camp_event(request):
    
    return render(request,'camp/campevent.html')

@login_required()
@group_required('po','vs')
def add_camp_event(request):

    camp={
        'camp':Camp.objects.all()
    }
    if request.method == "POST":
        event_name = request.POST.get('event_name')
        date = request.POST.get('date')
        description = request.POST.get('description')
        camp=request.POST.get('camp')
        camp=Camp.objects.get(camp_id=camp)
        camp_event = Camp_event(camp=camp,event_name=event_name, des=description,date=date)
        camp_event.save()
        message = "Submitted Successfully"
        return redirect(reverse('add_camp_event') + '?message=' + message)
    return render(request, 'camp/add_camp_event.html',camp)


@login_required()
@group_required('po','vs')
def campevent(request):
    return render(request,'camp/campuploads.html')

    

@login_required()
@group_required('po','vs')
def camp_photo(request):
    eve = {
        'even': Camp_event.objects.all().order_by('date').values()
    }
    if request.method == "POST" and request.FILES:
        photo=request.FILES.get('photo')
        event_id=request.POST.get('event_name')
        event_id=Camp_event.objects.get(event_id=event_id)
        if Camp_event_photos.objects.filter(event=event_id).count() >= 3:
            message = "Three images are already uploaded for this event."
            return redirect(reverse('camp_photo') + '?message1=' + message)
        else:
            ev=Camp_event_photos(photo=photo,event=event_id)
            ev.save()
            message = "Uploaded Sucessfully"
            return redirect(reverse('camp_photo') + '?message2=' + message)
    return render(request,'camp/camp_photo.html',eve)

    
    
@login_required()
@group_required('po','vs')
def view_camp_event(request):

    event = {
        'camp_events': Camp_event.objects.select_related('camp').all()
    }
    return render(request, 'camp/camp_view_event.html', event)
    
    
@login_required()
@group_required('po','vs')
def edit_camp_event(request,pk):

    eve={
        'eve':Camp_event.objects.filter(event_id=pk)
    }
    ev=Camp_event.objects.get(event_id=pk)
    event_Photos = Camp_event_photos.objects.filter(event=ev)
    event1 = get_object_or_404(Camp_event, pk=pk)
    if request.method=='POST':
        event_name=request.POST.get('event_name')
        date=request.POST.get('date')
        des=request.POST.get('des')
        event1.des=des
        event1.event_name=event_name
        event1.date=date
        event1.save()
        message = "Updated Successfully."
        return redirect(reverse('edit_camp_event', args=[pk]) + '?message=' + message)
    return render(request,'camp/edit_camp_event.html',eve)

@login_required()
@group_required('po','vs')
def delete_camp_images(request, pk, ev):

    photo = get_object_or_404(Camp_event_photos, id=pk)
    event_id = photo.event_id  
    photo.delete()
    
    photo_path = os.path.join(settings.MEDIA_ROOT, str(photo.photo))
    if os.path.exists(photo_path):
        os.remove(photo_path)
    
    message = "Image Deleted Successfully."
    return redirect(reverse('edit_camp_event', args=[event_id]) + '?message=' + message)


@login_required()
@group_required('po','vs')
def report_list_more_camp(request,pk):
    event={
        'eve':Camp_event.objects.filter(event_id=pk)
    }
    pics={
        'pics':Camp_event_photos.objects.filter(event_id=pk)
    }
    return render(request,'camp/report_list_more_camp.html',{**event,**pics})


@login_required()
@group_required('po','vs')
def delete_camp_event(request,pk):
    try:
        event = get_object_or_404(Camp_event, pk=pk)
        print(event)
        event.delete()
        return redirect('camp_event_view')
    except Exception:
        return render(request,'nss/error.html')

@login_required()
@group_required('po','vs')
def camp_report(request):
    report={
        'event':Camp_event.objects.all()
    }
    pics={
        'pics':Camp_event_photos.objects.all()
    }
    return render(request,'camp/camp_report.html',{**report,**pics})


def generate_event_report(request, event_id):
    # Get the event and related data
    event = Event.objects.get(event_id=event_id)
    event_details = Event_details.objects.filter(event=event)
    attendance_list = Attendance.objects.filter(event=event)
    event_photos = Event_Photos.objects.filter(event=event)
    
    # Create a new Document
    doc = Document()
    doc.add_heading(f'Report for {event.event_name}', 0)

    # Add event date
    doc.add_paragraph(f'{event.date}')

    # Add event details
    for detail in event_details:
        doc.add_paragraph(f'{detail.des}')

    # Add event photos
    for photo in event_photos:
        doc.add_picture(photo.photo.path, width=Inches(4))

    # Prepare the response to download the document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={event.event_name}_report.docx'
    
    # Save the document to the response
    doc.save(response)
    return response
